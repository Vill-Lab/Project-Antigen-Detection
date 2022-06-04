import cv2, os, sqlite3, re, shutil, datetime
import numpy as np
from pathlib import Path
from PIL import Image
from PIL import ImageFile
from docx import Document
import uuid
ImageFile.LOAD_TRUNCATED_IMAGES = True

def img_concat(temp_path, doc):
    path_list = temp_path.split('\n')
    img1 = cv_img_read(path_list[0])
    text = path_list[0]
    height, width = img1.shape[:-1]
    for p in path_list[1:]:
        text = text + ' ' * 4 + p
        t = cv_img_read(p)
        t_height, t_width = t.shape[:-1]
        factor = height / t_height
        t_width = int(t_width * factor)
        t = cv2.resize(t, (t_width, height), interpolation = cv2.INTER_AREA)
        img1 = cv2.hconcat([img1, t])
    img1 = cv2.resize(img1, (512, 512), interpolation = cv2.INTER_AREA)
    img1_path = Path(uuid.uuid4().hex + r'.png')
    cv2.imwrite(str(img1_path), img1)
    text = text + ' ' * 4 + '存在重复提交!'
    doc.add_paragraph(text)
    doc.add_picture(str(img1_path))
    img1_path.unlink()

def cv_img_read(path):
    #用PIL的路径
    img = Image.open(path).convert('RGB')
    #因为opencv读取是按照BGR的顺序，所以这里转换一下即可
    img_bgr = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
    return img_bgr

def main(im_name, cur, doc, repeatedLog, repeatedDir):
    img_path = str(im_name)
    inputImage = cv_img_read(img_path)
    detect_obj = cv2.wechat_qrcode_WeChatQRCode('detect.prototxt','detect.caffemodel','sr.prototxt','sr.caffemodel')
    res,_ = detect_obj.detectAndDecode(inputImage)
    num = len(res)
    if num > 0:
        for i in range(num):
            qrResult = res[i]
            #携带二维码的蓝色试剂
            if re.match( r'^http.*', qrResult):
                id = qrResult
            #携带二维码的白色试剂
            else:
                id = qrResult
            print(id)
            try:
                cur.execute(f"insert into Community (qrID, img_path) values ('{id}', '{img_path}')")
            except sqlite3.IntegrityError:
                with repeatedLog.open('a') as w:
                    w.write(im_name.name + "为重复提交!\n")
                cur.execute(f"select img_path from Community where qrID='{id}'")
                #import ipdb; ipdb.set_trace()
                temp_path = cur.fetchall()[0][0]
                temp_path = img_path + '\n' + temp_path
                img_concat(temp_path, doc)
                cur.execute(f"update Community set img_path='{temp_path}' where qrID='{id}'")
                shutil.copy(im_name, repeatedDir / im_name.name)
            else:
                pass
        return num, 1
    else:
        return 0, 0

if __name__ == "__main__":
    now = datetime.datetime.now()
    doc = Document()
    doc.add_picture('logo.png')
    doc.add_heading(str(now.year) + "年" + str(now.month) + "月" + str(now.day) + "日" + "杨浦区仁德居委67弄小区抗原数据监测报告", level=0)
    img_num= 0
    detected_qrcode = 0
    img_has_qrcode = 0
    p = Path('./img')
    repeatedDir = Path('./repeatedDir')
    if not repeatedDir.exists():
        repeatedDir.mkdir()
    else:
        shutil.rmtree(str(repeatedDir), ignore_errors=True)
        repeatedDir.mkdir()
    repeatedLog = Path('./repeatedLog.txt')
    if not repeatedLog.exists():
        repeatedLog.touch()
    else:
        repeatedLog.unlink()
        repeatedLog.touch()
    conn = sqlite3.connect('./community.db')
    cur = conn.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS Community
                    (qrID text primary key, img_path text)''')
    for e in p.glob("./5月1日/*.*g"): #p.glob("./*.*g"):
        img_num += 1
        try:
            t1, t2= main(e, cur, doc, repeatedLog, repeatedDir)
            detected_qrcode += t1
            img_has_qrcode += t2
        except OSError:
            continue
        except ValueError:
            continue
        else:
            continue
    cur.execute("select * from Community")
    print(cur.fetchall())
    print("二维码数量:", detected_qrcode)
    print("拥有二维码图片的数量:", img_has_qrcode)
    print("图片数量:", img_num)
    conn.commit()
    cur.close()
    conn.close()
    doc.save(str(now.year) + "年" + str(now.month) + "月" + str(now.day) + "日" + "杨浦区仁德居委67弄小区抗原数据监测报告" + r".docx")